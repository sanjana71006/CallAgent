import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

COLOR_NAVY = RGBColor(30, 58, 138)       # #1E3A8A
COLOR_SLATE = RGBColor(71, 85, 105)      # #475569
COLOR_TEXT = RGBColor(15, 23, 42)        # #0F172A
COLOR_GREEN = RGBColor(5, 150, 105)      # #059669
COLOR_AMBER = RGBColor(217, 119, 6)      # #D97706
COLOR_RED = RGBColor(220, 38, 38)        # #DC2626
HEX_NAVY = "1E3A8A"
HEX_ZEBRA = "F8FAFC"
HEX_BORDER = "CBD5E1"

def init_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    return doc

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="6"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_title_page(doc):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("CALLMATE AI")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(32)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_NAVY

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(20)
    p_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Complete Project Architecture, Engineering & Technical Documentation")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_SLATE

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "EFF6FF")
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="3B82F6"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="3B82F6"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="3B82F6"/>'
        f'<w:right w:val="single" w:sz="12" w:space="0" w:color="3B82F6"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p_box = cell.paragraphs[0]
    p_box.paragraph_format.space_after = Pt(4)
    p_box.paragraph_format.line_spacing = 1.2
    
    p_box.add_run("Project Name: ").bold = True
    p_box.add_run("CallMate AI — Privacy-First Intelligent Call Assistant\n")
    p_box.add_run("GitHub Repository: ").bold = True
    p_box.add_run("https://github.com/sanjana71006/CallAgent\n")
    p_box.add_run("150-Char Synopsis: ").bold = True
    p_box.add_run("CallMate AI is a privacy-first call assistant that screens unknown calls, transcribes voice live, blocks spam, and guides couriers with Google Gemini.\n")
    p_box.add_run("Version: ").bold = True
    p_box.add_run("v1.0.0 (Release Build 2026081841 - CallMate_AI_v1.0.0_debug.apk)\n")
    p_box.add_run("Documentation Date: ").bold = True
    p_box.add_run("August 31, 2026\n")
    p_box.add_run("Architecture Scope: ").bold = True
    p_box.add_run("Android Client (Kotlin/Compose/Room) + Cloud Backend (Node.js/Express/MongoDB Atlas) + AI Engine (FastAPI/Google Gemini 3.6 Flash)\n")
    p_box.add_run("Source of Truth: ").bold = True
    p_box.add_run("Verified Live Codebase Implementation (No Invented Features)")

    doc.add_paragraph().paragraph_format.space_after = Pt(28)

    p_abs_title = doc.add_paragraph()
    p_abs_title.paragraph_format.space_before = Pt(8)
    p_abs_title.paragraph_format.space_after = Pt(4)
    r_at = p_abs_title.add_run("DOCUMENT ABSTRACT & EXECUTIVE SUMMARY")
    r_at.bold = True
    r_at.font.color.rgb = COLOR_NAVY
    r_at.font.size = Pt(12)

    p_abs = doc.add_paragraph(
        "This document provides a factual, exhaustive, and rigorously verified technical record of the CallMate AI system. "
        "CallMate AI is an intelligent call management and screening platform engineered with an offline-first Android client (Kotlin, Jetpack Compose, Room + SQLite), "
        "a cloud authentication, crowdsourced spam, and address database (Node.js, Express, MongoDB Atlas), and an AI call screening inference server (Python FastAPI with Google Gemini 3.6 Flash integration). "
        "Every section in this document describes the precise state of the codebase, strictly differentiating implemented features from partially implemented modules and planned future work."
    )
    p_abs.paragraph_format.line_spacing = 1.15

    doc.add_page_break()

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_SLATE
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEXT
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEXT
    return p

def add_bullet(doc, prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.color.rgb = COLOR_TEXT
    r2 = p.add_run(" " + text)
    r2.font.color.rgb = COLOR_TEXT
    return p

def add_callout(doc, text, title="NOTE:", status="info"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    bg_hex = "F1F5F9"
    border_color = "3B82F6"
    title_color = COLOR_NAVY
    if status == "warning":
        bg_hex = "FFFBEB"
        border_color = "F59E0B"
        title_color = COLOR_AMBER
    elif status == "success":
        bg_hex = "ECFDF5"
        border_color = "10B981"
        title_color = COLOR_GREEN
    elif status == "danger":
        bg_hex = "FEF2F2"
        border_color = "EF4444"
        title_color = COLOR_RED

    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = title_color
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "0F172A")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(241, 245, 249)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_table(doc, col_widths, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table, color="CBD5E1", sz="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = Inches(col_widths[i])
        set_cell_background(hdr_cells[i], HEX_NAVY)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        trPr_row = row._tr.get_or_add_trPr()
        trPr_row.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        bg_color = HEX_ZEBRA if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = Inches(col_widths[col_idx])
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            if text in ["IMPLEMENTED", "ACTIVE", "READY", "YES"]:
                run.bold = True
                run.font.color.rgb = COLOR_GREEN
            elif text in ["PARTIALLY IMPLEMENTED", "WARNING", "DEGRADED", "PARTIAL"]:
                run.bold = True
                run.font.color.rgb = COLOR_AMBER
            elif text in ["PLANNED / NOT YET IMPLEMENTED", "NOT IMPLEMENTED", "NO", "PLANNED"]:
                run.bold = True
                run.font.color.rgb = COLOR_RED
            else:
                run.font.color.rgb = COLOR_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
